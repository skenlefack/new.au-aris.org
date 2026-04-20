#!/usr/bin/env python3
"""
Generate Word report for Historical Data cleanup operations.
"""
from __future__ import annotations

from datetime import date
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUTPUT = "C:/new.au-aris.org/docs/ARIS4_Rapport_Nettoyage_Donnees_Historiques.docx"


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=9, color=None):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def shade_row(row, hex_color):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = tcPr.makeelement(qn("w:shd"), {
            qn("w:fill"): hex_color,
            qn("w:val"): "clear",
        })
        tcPr.append(shading)


def add_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = table.rows[0]
    shade_row(hdr, header_color)
    for i, h in enumerate(headers):
        set_cell(hdr.cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                 size=9, color=RGBColor(255, 255, 255))

    # Data
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        if ri % 2 == 1:
            shade_row(row, "F2F7FB")
        for ci, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.RIGHT if isinstance(val, (int, float)) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell(row.cells[ci], f"{val:,}" if isinstance(val, int) else str(val), align=align)

    if col_widths:
        for ri in range(len(table.rows)):
            for ci, w in enumerate(col_widths):
                table.rows[ri].cells[ci].width = Cm(w)

    return table


def main():
    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── TITLE ────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AFRICAN UNION — INTER-AFRICAN BUREAU FOR ANIMAL RESOURCES")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ARIS 4.0 — Animal Resources Information System")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph()

    title = doc.add_heading("Rapport de nettoyage des donnees historiques", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Environnement : Staging (test.au-aris.org)  —  Date : {date.today().strftime('%d/%m/%Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # ── 1. CONTEXTE ──────────────────────────────────────────────────
    doc.add_heading("1. Contexte", level=2)
    doc.add_paragraph(
        "Les donnees historiques d'ARIS 3 (2008-2025) ont ete importees dans le module "
        "Historical Data du datalake ARIS 4 sur l'environnement de staging. "
        "L'import initial avait produit 15 datasets totalisant 809 460 enregistrements "
        "couvrant 46 pays africains."
    )
    doc.add_paragraph(
        "Deux operations de nettoyage ont ete effectuees :"
    )
    bullets = [
        "Renommage et consolidation des datasets (suppression du prefixe « ARIS 3 — », "
        "fusion des 9 parties du Monthly Animal Health Reports en un seul dataset)",
        "Normalisation de la colonne « disease » pour harmoniser les noms de maladies "
        "vers les identifiants UUID de la table de reference animal_health.ref_diseases",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # ── 2. RENOMMAGE DES DATASETS ────────────────────────────────────
    doc.add_heading("2. Renommage et consolidation des datasets", level=2)

    doc.add_heading("2.1 Renommage (suppression du prefixe « ARIS 3 — »)", level=3)
    doc.add_paragraph(
        "Les 6 datasets individuels ont ete renommes pour supprimer la mention "
        "« ARIS 3 — » au debut du nom, produisant des noms plus clairs et concis :"
    )

    rename_rows = [
        ("ARIS 3 — Emergency Disease Reports (2008-2025)",     "Emergency Disease Reports (2008-2025)"),
        ("ARIS 3 — Mass Vaccination Campaigns (2008-2025)",    "Mass Vaccination Campaigns (2008-2025)"),
        ("ARIS 3 — Meat Inspection Reports (2015-2018)",       "Meat Inspection Reports (2015-2018)"),
        ("ARIS 3 — Monthly Vaccination Reports (2008-2025)",   "Monthly Vaccination Reports (2008-2025)"),
        ("ARIS 3 — Animal Population & Composition (2008-2025)", "Animal Population & Composition (2008-2025)"),
        ("ARIS 3 — Animal Import & Export Records (2014-2017)", "Animal Import & Export Records (2014-2017)"),
    ]
    add_table(doc, ["Ancien nom", "Nouveau nom"], rename_rows, col_widths=[8.5, 8.5])

    doc.add_heading("2.2 Fusion du Monthly Animal Health Reports", level=3)
    doc.add_paragraph(
        "Le fichier Monthly Animal Health Reports contenait 801 318 lignes et avait ete "
        "importe en 9 parties (chunks de ~100 000 lignes) pour contourner les limites "
        "memoire de l'API d'import. Ces 9 datasets ont ete fusionnes en un seul :"
    )

    parts_rows = [
        ("Part 1/9", 100000),
        ("Part 2/9", 100000),
        ("Part 3/9", 100000),
        ("Part 4/9", 100000),
        ("Part 5/9", 100000),
        ("Part 6/9", 100000),
        ("Part 7/9", 100000),
        ("Part 8/9", 100000),
        ("Part 9/9", 1318),
        ("TOTAL fusionne", 801318),
    ]
    add_table(doc, ["Partie", "Lignes"], parts_rows, col_widths=[10, 7])

    doc.add_paragraph()
    doc.add_paragraph(
        "Operation SQL : CREATE TABLE ... (LIKE part1 INCLUDING ALL), puis INSERT INTO ... "
        "SELECT depuis chaque partie, suppression des 9 anciens datasets et tables."
    )

    doc.add_heading("2.3 Etat final des datasets", level=3)
    final_ds = [
        ("Monthly Animal Health Reports (2008-2025)", "animal_health", 801318),
        ("Emergency Disease Reports (2008-2025)",     "animal_health", 128),
        ("Mass Vaccination Campaigns (2008-2025)",    "animal_health", 375),
        ("Meat Inspection Reports (2015-2018)",       "animal_health", 992),
        ("Monthly Vaccination Reports (2008-2025)",   "animal_health", 6308),
        ("Animal Population & Composition (2008-2025)", "livestock",  249),
        ("Animal Import & Export Records (2014-2017)", "trade_sps",   90),
    ]
    add_table(doc, ["Dataset", "Domaine", "Lignes"],
              final_ds, col_widths=[9, 3.5, 3])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Resultat : 15 datasets → 7 datasets. Total : 809 460 enregistrements preserves.")
    run.bold = True

    # ── 3. NORMALISATION DES MALADIES ────────────────────────────────
    doc.add_heading("3. Normalisation de la colonne « disease »", level=2)

    doc.add_heading("3.1 Probleme identifie", level=3)
    doc.add_paragraph(
        "La colonne « disease » du dataset Monthly Animal Health Reports contenait "
        "un melange heterogene de :"
    )
    issue_bullets = [
        "UUID de ref_diseases (ex: 30000000-0000-4000-b000-000000000008 pour Rabies) — 187 442 lignes",
        "Noms textuels en anglais MAJUSCULES (ex: RABIES, ANTHRAX, BLACKLEG) — majorite des lignes textuelles",
        "Noms en francais (ex: Rage, Fievre Aphteuse, Dermatose Nodulaire)",
        "Variantes avec fautes de frappe (ex: Trypanosomiasi, Typanosomosis, Lumbyskin)",
        "Abreviations/codes (ex: FMD, PPR, LSD, ND, CCPP)",
        "Valeurs vides (aucune maladie renseignee) — 294 121 lignes",
        "Bruit (noms de localites, en-tetes de formulaire, mesures de controle) — ~29 000 lignes",
    ]
    for b in issue_bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3.2 Approche", level=3)
    doc.add_paragraph(
        "Un dictionnaire de correspondance de 254 variantes textuelles couvrant 54 maladies "
        "a ete construit a partir de :"
    )
    approach_bullets = [
        "Les noms multilingues (EN, FR, PT, AR) de la table animal_health.ref_diseases (88 maladies)",
        "Les codes officiels (FMD, CBPP, PPR, ASF, HPAI, etc.)",
        "Les variantes, abreviations et fautes de frappe trouvees dans les donnees",
        "Matching case-insensitive avec remplacement underscore/espace",
    ]
    for b in approach_bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph(
        "Chaque variante textuelle a ete remplacee par l'UUID correspondant de ref_diseases "
        "via des requetes UPDATE SQL groupees par maladie."
    )

    doc.add_heading("3.3 Resultats de la normalisation", level=3)

    # Before / After
    doc.add_paragraph("Etat avant normalisation :", style="List Bullet")
    before_rows = [
        ("UUID (deja normalise)",  187442, "23.4%"),
        ("Texte (noms de maladies)", 319755, "39.9%"),
        ("Vide (pas de maladie)",   294121, "36.7%"),
        ("Total",                    801318, "100%"),
    ]
    add_table(doc, ["Categorie", "Lignes", "%"], before_rows, col_widths=[7, 4, 3])

    doc.add_paragraph()
    doc.add_paragraph("Etat apres normalisation :", style="List Bullet")
    after_rows = [
        ("UUID (normalise)",    477510, "59.6%"),
        ("Vide (pas de maladie)", 294121, "36.7%"),
        ("Texte residuel (non-maladies)", 29687, "3.7%"),
        ("Total",                801318, "100%"),
    ]
    add_table(doc, ["Categorie", "Lignes", "%"], after_rows, col_widths=[7, 4, 3])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("290 068 lignes normalisees (90.7% des lignes textuelles).")
    run.bold = True

    doc.add_heading("3.4 Top 20 des maladies normalisees", level=3)
    top20 = [
        ("Rabies",                          "...008", 64691, 8),
        ("Anthrax",                         "...009", 27499, 1),
        ("Blackleg",                        "...015", 20241, 9),
        ("Anaplasmosis",                    "...021", 18569, 4),
        ("Brucellosis",                     "...010", 16305, 10),
        ("Babesiosis",                      "...022", 13962, 5),
        ("Fasciolosis",                     "...023", 12802, 4),
        ("Heartwater",                      "...020", 10298, 2),
        ("Haemorrhagic Septicaemia",        "...042", 8728,  7),
        ("Mange",                           "...024", 8378,  7),
        ("Lumpy Skin Disease",              "...007", 8205,  6),
        ("Newcastle Disease",               "...012", 8041,  13),
        ("Sheep Pox",                       "...016", 6311,  6),
        ("Leptospirosis",                   "...060", 5909,  2),
        ("Infectious Bursal Disease",       "...054", 5684,  3),
        ("Contagious Agalactia",            "...045", 5665,  2),
        ("Mastitis",                        "...025", 5172,  2),
        ("Infectious Coryza",              "...074", 5081,  6),
        ("Foot and Mouth Disease",          "...001", 3802,  10),
        ("Fowl Pox",                        "...058", 3549,  12),
    ]
    add_table(doc, ["Maladie", "UUID (suffixe)", "Lignes mises a jour", "Variantes"],
              top20, col_widths=[5.5, 3, 4, 2.5])

    doc.add_heading("3.5 Texte residuel (non normalise)", level=3)
    doc.add_paragraph(
        "Les 29 687 lignes textuelles restantes (3.7%) ne correspondent pas a des maladies "
        "de la table de reference. Elles comprennent :"
    )
    residual_bullets = [
        "Maladies sans entree dans ref_diseases : DERMATOPHILOSIS (6 253), Pasteurellose (various), "
        "Canine Parvoviral Enteritis, Toxoplasmosis, Swine erysipelas, etc.",
        "Noms de localites mal places dans la colonne disease : North Bungoma, Kieini West/East, "
        "Bumula, Matungu, Kuria West, Butula, etc.",
        "En-tetes de formulaire : « Reporting Unit: », « Disease », « Month », « Date Report: »",
        "Mesures de controle : « Treatment », « Vaccination », « Stamping out »",
        "Valeurs non exploitables : « Unknown disease », « NA », codes internes (DJSG1309, etc.)",
    ]
    for b in residual_bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph(
        "Ces lignes restent en l'etat dans la base. Les maladies manquantes dans ref_diseases "
        "(Dermatophilosis, Pasteurellosis, Bovine Ephemeral Fever, etc.) pourront etre ajoutees "
        "ulterieurement au referentiel si necessaire."
    )

    # ── 4. MODIFICATIONS DU CODE ─────────────────────────────────────
    doc.add_heading("4. Modifications du code source", level=2)
    doc.add_paragraph(
        "En plus des operations sur la base de donnees, les fichiers suivants ont ete mis a jour "
        "pour supprimer les references a « ARIS 3 » dans l'interface utilisateur :"
    )

    code_rows = [
        ("apps/web/src/components/historical/HistoricalDataSection.tsx",
         "Suppression du prefixe « ARIS 3 » dans l'affichage des datasets"),
        ("apps/web/src/app/(dashboard)/animal-health/page.tsx",
         "Titre « Historical Data (2008-2025) » sans mention ARIS 3"),
        ("apps/web/src/app/(dashboard)/livestock/page.tsx",
         "Idem pour la page Livestock"),
        ("apps/web/src/app/(dashboard)/trade/page.tsx",
         "Idem pour la page Trade & SPS"),
        ("apps/web/src/app/(dashboard)/historical/dashboard/page.tsx",
         "Sous-titre du dashboard analytique"),
        ("apps/web/src/messages/en.json",
         "Traduction EN du sous-titre dashboard"),
        ("apps/web/src/messages/fr.json",
         "Traduction FR du sous-titre dashboard"),
        ("deploy/scripts/_import_historical_to_aris4.py",
         "Noms des datasets pour les futurs imports"),
        ("deploy/scripts/_import_historical_chunked.py",
         "Noms des chunks pour les futurs imports"),
        ("deploy/scripts/_create_superset_dashboard.py",
         "Description des vues Superset"),
        ("deploy/scripts/_setup_superset_historical.py",
         "Description des saved queries Superset"),
    ]
    add_table(doc, ["Fichier", "Modification"], code_rows, col_widths=[8.5, 8.5])

    # ── 5. SCRIPTS ───────────────────────────────────────────────────
    doc.add_heading("5. Scripts utilises", level=2)
    scripts = [
        ("_rename_historical_datasets.py",
         "Renommage des 6 datasets + fusion des 9 parties Monthly Animal Health Reports"),
        ("_normalize_historical_diseases.py",
         "Normalisation de la colonne disease (254 variantes → 54 UUIDs ref_diseases)"),
    ]
    add_table(doc, ["Script", "Description"], scripts, col_widths=[7, 10])
    doc.add_paragraph()
    doc.add_paragraph(
        "Les deux scripts supportent les options --env stg|prod et --dry-run. "
        "Ils se connectent via SSH/paramiko au serveur de base de donnees et executent "
        "les requetes SQL directement dans le conteneur PostgreSQL."
    )

    # ── 6. PROCHAINES ETAPES ─────────────────────────────────────────
    doc.add_heading("6. Prochaines etapes", level=2)
    next_steps = [
        "Deployer le frontend (web) sur staging pour refleter les changements d'affichage",
        "Enrichir la table ref_diseases avec les maladies manquantes (Dermatophilosis, "
        "Pasteurellosis, Bovine Ephemeral Fever, etc.) pour normaliser les lignes residuelles",
        "Appliquer les memes operations sur l'environnement de production (--env prod)",
        "Verifier les dashboards Superset avec les nouvelles donnees normalisees",
    ]
    for i, s in enumerate(next_steps, 1):
        doc.add_paragraph(s, style="List Number")

    # ── Footer ───────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— Fin du rapport —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Genere le {date.today().strftime('%d/%m/%Y')} — ARIS 4.0 / AU-IBAR")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(OUTPUT)
    print(f"[done] Report saved to {OUTPUT}")


if __name__ == "__main__":
    main()
