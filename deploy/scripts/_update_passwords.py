#!/usr/bin/env python3
"""
Update ALL user passwords on PROD and STG to a new password.
Generates bcrypt hash in Python, writes SQL file, SFTPs to DB VM,
executes via psql -f to avoid shell escaping issues with $ in bcrypt hashes.
Also generates a Word document with user credentials table.
"""

import paramiko
import bcrypt
import sys
import time

# ── Configuration ──────────────────────────────────────────────────────────
NEW_PASSWORD = "Aris2026@@4!0"
SSH_USER = "arisadmin"
SSH_PASS = "@u-1baR.0rg$U24"

ENVIRONMENTS = {
    "PROD": {
        "db_host": "10.202.101.185",
        "pg_container": "aris-postgres",
        "db_name": "aris",
        "db_user": "aris",
    },
    "STG": {
        "db_host": "10.202.101.148",
        "pg_container": "aris-stg-postgres",
        "db_name": "aris",
        "db_user": "aris",
    },
}


def ssh_connect(host: str) -> paramiko.SSHClient:
    client = paramiko.SSHClient()
    client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    client.connect(host, username=SSH_USER, password=SSH_PASS, timeout=15)
    return client


def sudo(ssh: paramiko.SSHClient, cmd: str, timeout: int = 120) -> str:
    """Run a sudo command via SSH using the proven channel + sendall pattern."""
    ch = ssh.get_transport().open_session()
    ch.settimeout(timeout)
    ch.exec_command("sudo -S " + cmd)
    ch.sendall((SSH_PASS + "\n").encode())
    time.sleep(0.5)
    out = b""
    while ch.recv_ready() or not ch.exit_status_ready():
        if ch.recv_ready():
            out += ch.recv(65536)
        else:
            time.sleep(0.3)
            if ch.exit_status_ready() and not ch.recv_ready():
                break
    lines = [l for l in out.decode("utf-8", "replace").splitlines()
             if "[sudo]" not in l]
    return "\n".join(lines)


def run_cmd(ssh: paramiko.SSHClient, cmd: str, timeout: int = 30) -> str:
    """Run a non-sudo command via SSH."""
    stdin, stdout, stderr = ssh.exec_command(cmd, timeout=timeout)
    out = stdout.read().decode("utf-8", errors="replace")
    return out.strip()


def generate_bcrypt_hash(password: str) -> str:
    """Generate bcrypt hash locally in Python."""
    salt = bcrypt.gensalt(rounds=12)
    hashed = bcrypt.hashpw(password.encode("utf-8"), salt)
    return hashed.decode("utf-8")


def update_passwords_on_env(env_name: str, config: dict, bcrypt_hash: str) -> list:
    """Update all user passwords on a given environment. Returns list of users."""
    print(f"\n{'='*60}")
    print(f"  {env_name} — DB VM: {config['db_host']}")
    print(f"{'='*60}")

    client = ssh_connect(config["db_host"])

    # 1. Create SQL file content — use dollar-quoting to avoid ANY escaping issues
    sql_content = f"""-- Update all user passwords
BEGIN;
UPDATE users SET password_hash = $bcrypt${bcrypt_hash}$bcrypt$;
COMMIT;
SELECT email, role, is_active FROM users ORDER BY role, email;
"""

    # 2. SFTP the SQL file to the DB VM
    print("  [1/4] Uploading SQL file via SFTP...")
    sftp = client.open_sftp()
    with sftp.file("/tmp/update_passwords.sql", "w") as f:
        f.write(sql_content)
    sftp.close()
    print("        Done.")

    # 3. Copy SQL file into postgres container
    print("  [2/4] Copying SQL into postgres container...")
    out = sudo(client, f"docker cp /tmp/update_passwords.sql {config['pg_container']}:/tmp/update_passwords.sql")
    if out.strip():
        print(f"        {out.strip()}")
    print("        Done.")

    # 4. Execute SQL via psql -f (no shell escaping needed!)
    print("  [3/4] Executing password update...")
    result = sudo(
        client,
        f"docker exec {config['pg_container']} psql -U {config['db_user']} -d {config['db_name']} -f /tmp/update_passwords.sql",
        timeout=30,
    )
    print(f"        Result:\n{result}")

    # 5. Get user list with pipe-separated output for parsing
    print("  [4/4] Getting user list...")
    verify = sudo(
        client,
        f"docker exec {config['pg_container']} psql -U {config['db_user']} -d {config['db_name']} -t -A -F'|' -c \"SELECT email, role, is_active FROM users ORDER BY role, email;\"",
        timeout=15,
    )

    users = []
    for line in verify.strip().split("\n"):
        parts = line.strip().split("|")
        if len(parts) >= 3:
            users.append({
                "email": parts[0].strip(),
                "role": parts[1].strip(),
                "is_active": parts[2].strip(),
            })

    print(f"        Found {len(users)} users.")

    # Cleanup
    run_cmd(client, "rm -f /tmp/update_passwords.sql")
    sudo(client, f"docker exec {config['pg_container']} rm -f /tmp/update_passwords.sql")

    client.close()
    return users


def verify_login(env_name: str, app_host: str, email: str, password: str) -> bool:
    """Test login via the credential service API."""
    client = ssh_connect(app_host)
    cmd = (
        f'curl -s -o /dev/null -w "%{{http_code}}" '
        f'-X POST http://localhost:3002/api/v1/credential/auth/login '
        f'-H "Content-Type: application/json" '
        f"-d '{{\"email\":\"{email}\",\"password\":\"{password}\"}}'"
    )
    result = run_cmd(client, cmd, timeout=10)
    client.close()
    status = result.strip()
    print(f"  Login test ({env_name}): {email} -> HTTP {status}")
    return status == "200" or status == "201"


def generate_word_doc(users: list, password: str, output_path: str):
    """Generate a Word document with user credentials table."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading("ARIS 4.0 — User Credentials", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Password: {password}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)

    doc.add_paragraph(f"Total users: {len(users)}")
    doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph()

    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["#", "Email", "Role", "Active"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for r in paragraph.runs:
                r.bold = True
                r.font.size = Pt(10)

    # Data rows
    for idx, user in enumerate(users, 1):
        row = table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = user["email"]
        row.cells[2].text = user["role"]
        row.cells[3].text = "Yes" if user["is_active"] in ("t", "true", "True", True) else "No"

        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for r in paragraph.runs:
                    r.font.size = Pt(9)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(1)
        row.cells[1].width = Cm(8)
        row.cells[2].width = Cm(5)
        row.cells[3].width = Cm(2)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("CONFIDENTIAL — AU-IBAR ARIS 4.0")
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    print(f"\n  Word document saved: {output_path}")


def main():
    print("=" * 60)
    print("  ARIS 4.0 — Password Update Script")
    print(f"  New password: {NEW_PASSWORD}")
    print("=" * 60)

    # Step 1: Generate bcrypt hash locally
    print("\n  Generating bcrypt hash...")
    bcrypt_hash = generate_bcrypt_hash(NEW_PASSWORD)
    print(f"  Hash: {bcrypt_hash}")
    print(f"  Hash length: {len(bcrypt_hash)} chars")

    # Step 2: Update passwords on each environment
    all_users = {}
    for env_name, config in ENVIRONMENTS.items():
        try:
            users = update_passwords_on_env(env_name, config, bcrypt_hash)
            all_users[env_name] = users
        except Exception as e:
            print(f"\n  ERROR on {env_name}: {e}")
            import traceback
            traceback.print_exc()

    # Step 3: Verify login on each environment
    print(f"\n{'='*60}")
    print("  Verifying logins...")
    print(f"{'='*60}")

    app_hosts = {"PROD": "10.202.101.183", "STG": "10.202.101.146"}
    for env_name in all_users:
        verify_login(env_name, app_hosts[env_name], "admin@au-aris.org", NEW_PASSWORD)

    # Step 4: Generate Word document
    print(f"\n{'='*60}")
    print("  Generating Word document...")
    print(f"{'='*60}")

    # Use PROD users for the document (should be same as STG)
    doc_users = all_users.get("PROD", all_users.get("STG", []))
    if doc_users:
        output_path = r"C:\new.au-aris.org\ARIS_User_Credentials.docx"
        try:
            generate_word_doc(doc_users, NEW_PASSWORD, output_path)
        except ImportError:
            print("\n  python-docx not installed. Installing...")
            import subprocess
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            generate_word_doc(doc_users, NEW_PASSWORD, output_path)

    print(f"\n{'='*60}")
    print("  DONE!")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
